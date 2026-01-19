#!/usr/bin/env python3
"""
Create Word documents for each chapter with all questions and answers.
"""

import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Wine color RGB
WINE_COLOR = RGBColor(114, 47, 55)  # Dark wine/burgundy color

# Base path for questions
BASE_PATH = "/Users/ramsabode/vibe-coding/ani-questions-vercel/questions_cache"
OUTPUT_PATH = "/Users/ramsabode/vibe-coding/ani-questions-vercel/word_documents"

# Question type display names
QUESTION_TYPE_NAMES = {
    'textbook_qa': 'Textbook Questions & Answers',
    'mcq': 'Multiple Choice Questions',
    'fill_in_blanks': 'Fill in the Blanks',
    'true_false': 'True or False',
    'match_the_following': 'Match the Following',
    'short_answer': 'Short Answer Questions',
    'long_answer': 'Long Answer Questions',
    'assertion_reason': 'Assertion and Reason',
    'give_reason': 'Give Reasons',
    'differentiate_between': 'Differentiate Between',
    'numericals': 'Numerical Problems',
    'hots': 'Higher Order Thinking Skills (HOTS)',
    'case_study': 'Case Study Questions',
    'name_the_following': 'Name the Following',
    'tricky_questions': 'Tricky Questions'
}

def create_document():
    """Create a new document with wine-colored heading styles."""
    doc = Document()

    # Modify heading styles to use wine color
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        style.font.color.rgb = WINE_COLOR
        style.font.bold = True
        if i == 1:
            style.font.size = Pt(18)
        elif i == 2:
            style.font.size = Pt(14)
        else:
            style.font.size = Pt(12)

    return doc

def add_chapter_title(doc, subject, chapter):
    """Add chapter title to document."""
    title = doc.add_heading(f"{subject} - {chapter}", level=0)
    title.runs[0].font.color.rgb = WINE_COLOR
    title.runs[0].font.size = Pt(24)
    doc.add_paragraph()

def add_section_title(doc, title):
    """Add section title with wine color."""
    heading = doc.add_heading(title, level=1)
    return heading

def add_question_type_title(doc, title):
    """Add question type title."""
    heading = doc.add_heading(title, level=2)
    return heading

def format_question_number(doc, num, question_text):
    """Add a question with number."""
    para = doc.add_paragraph()
    run = para.add_run(f"Q{num}. ")
    run.bold = True
    run.font.color.rgb = WINE_COLOR
    para.add_run(question_text)
    return para

def format_answer(doc, answer_text, label="Answer:"):
    """Add an answer."""
    para = doc.add_paragraph()
    run = para.add_run(f"{label} ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
    para.add_run(str(answer_text))
    return para

def add_mcq_question(doc, num, q):
    """Add MCQ question with options and answer."""
    format_question_number(doc, num, q.get('question', ''))

    options = q.get('options', [])
    for i, opt in enumerate(options):
        letter = chr(65 + i)  # A, B, C, D
        doc.add_paragraph(f"   {letter}. {opt}")

    correct = q.get('correct_answer', '')
    explanation = q.get('explanation', '')

    format_answer(doc, correct)
    if explanation:
        para = doc.add_paragraph()
        run = para.add_run("Explanation: ")
        run.italic = True
        para.add_run(explanation)

    doc.add_paragraph()

def add_true_false_question(doc, num, q):
    """Add True/False question."""
    statement = q.get('statement', q.get('question', ''))
    format_question_number(doc, num, statement)

    answer = q.get('answer', q.get('correct_answer', ''))
    if answer == True or answer == 'True':
        answer = 'True'
    else:
        answer = 'False'

    explanation = q.get('explanation', '')

    format_answer(doc, answer)
    if explanation:
        para = doc.add_paragraph()
        run = para.add_run("Explanation: ")
        run.italic = True
        para.add_run(explanation)

    doc.add_paragraph()

def add_fill_in_blanks_question(doc, num, q):
    """Add Fill in the Blanks question."""
    format_question_number(doc, num, q.get('question', ''))

    answer = q.get('answer', q.get('correct_answer', ''))
    explanation = q.get('explanation', '')

    format_answer(doc, answer)
    if explanation:
        para = doc.add_paragraph()
        run = para.add_run("Explanation: ")
        run.italic = True
        para.add_run(explanation)

    doc.add_paragraph()

def add_match_question(doc, num, q):
    """Add Match the Following question."""
    instruction = q.get('instruction', q.get('question', 'Match the following:'))
    format_question_number(doc, num, instruction)

    pairs = q.get('pairs', [])

    # Show columns
    doc.add_paragraph("Column A" + " " * 30 + "Column B")
    for i, pair in enumerate(pairs):
        left = pair.get('left', '')
        right = pair.get('right', '')
        doc.add_paragraph(f"   {i+1}. {left}")

    doc.add_paragraph()
    for i, pair in enumerate(pairs):
        right = pair.get('right', '')
        doc.add_paragraph(f"   {chr(65+i)}. {right}")

    # Answer
    para = doc.add_paragraph()
    run = para.add_run("Answer: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    matches = []
    for i, pair in enumerate(pairs):
        matches.append(f"{i+1} - {pair.get('right', '')}")
    para.add_run(", ".join(matches))

    doc.add_paragraph()

def add_short_answer_question(doc, num, q):
    """Add Short Answer question."""
    format_question_number(doc, num, q.get('question', ''))

    answer = q.get('answer', q.get('correct_answer', ''))
    format_answer(doc, answer)

    keywords = q.get('keywords', [])
    if keywords:
        para = doc.add_paragraph()
        run = para.add_run("Keywords: ")
        run.italic = True
        para.add_run(", ".join(keywords))

    doc.add_paragraph()

def add_textbook_qa_question(doc, num, q):
    """Add Textbook Q&A question."""
    format_question_number(doc, num, q.get('question', ''))

    answer = q.get('answer', q.get('correct_answer', ''))
    format_answer(doc, answer)

    doc.add_paragraph()

def add_differentiate_question(doc, num, q):
    """Add Differentiate Between question."""
    term1 = q.get('term1', q.get('concept_a', 'Term 1'))
    term2 = q.get('term2', q.get('concept_b', 'Term 2'))

    format_question_number(doc, num, f"Differentiate between {term1} and {term2}")

    differences = q.get('differences', [])

    para = doc.add_paragraph()
    run = para.add_run("Answer:")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    for diff in differences:
        aspect = diff.get('aspect', '')
        point_a = diff.get('term1_point', diff.get('concept_a_point', ''))
        point_b = diff.get('term2_point', diff.get('concept_b_point', ''))

        if aspect:
            doc.add_paragraph(f"   {aspect}:")
        doc.add_paragraph(f"   • {term1}: {point_a}")
        doc.add_paragraph(f"   • {term2}: {point_b}")

    doc.add_paragraph()

def add_give_reason_question(doc, num, q):
    """Add Give Reason question."""
    statement = q.get('statement', q.get('question', ''))
    format_question_number(doc, num, f"Give reason: {statement}")

    reason = q.get('reason', q.get('answer', q.get('correct_answer', '')))
    format_answer(doc, reason)

    doc.add_paragraph()

def add_numerical_question(doc, num, q):
    """Add Numerical question."""
    format_question_number(doc, num, q.get('question', ''))

    given_data = q.get('given_data', [])
    if given_data:
        para = doc.add_paragraph()
        run = para.add_run("Given: ")
        run.bold = True
        para.add_run(", ".join(given_data))

    formula = q.get('formula', '')
    if formula:
        para = doc.add_paragraph()
        run = para.add_run("Formula: ")
        run.bold = True
        para.add_run(formula)

    steps = q.get('solution_steps', [])
    if steps:
        para = doc.add_paragraph()
        run = para.add_run("Solution:")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0)
        for i, step in enumerate(steps):
            doc.add_paragraph(f"   Step {i+1}: {step}")

    answer = q.get('answer', q.get('correct_answer', ''))
    if answer:
        format_answer(doc, answer, "Final Answer:")

    doc.add_paragraph()

def add_assertion_reason_question(doc, num, q):
    """Add Assertion Reason question."""
    assertion = q.get('assertion', '')
    reason = q.get('reason', '')

    para = doc.add_paragraph()
    run = para.add_run(f"Q{num}. ")
    run.bold = True
    run.font.color.rgb = WINE_COLOR

    doc.add_paragraph(f"Assertion (A): {assertion}")
    doc.add_paragraph(f"Reason (R): {reason}")

    options = q.get('options', [])
    for opt in options:
        doc.add_paragraph(f"   {opt}")

    correct = q.get('correct_answer', q.get('correct_option', ''))
    explanation = q.get('explanation', '')

    format_answer(doc, correct)
    if explanation:
        para = doc.add_paragraph()
        run = para.add_run("Explanation: ")
        run.italic = True
        para.add_run(explanation)

    doc.add_paragraph()

def add_case_study_question(doc, num, q):
    """Add Case Study question."""
    case_text = q.get('case_text', q.get('question', ''))
    format_question_number(doc, num, "Read the passage and answer the questions:")

    para = doc.add_paragraph()
    para.add_run(case_text).italic = True

    sub_questions = q.get('questions', [])
    for i, sq in enumerate(sub_questions):
        letter = chr(97 + i)  # a, b, c, d
        doc.add_paragraph(f"   ({letter}) {sq.get('question', '')}")

        para = doc.add_paragraph()
        run = para.add_run(f"   Answer: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0)
        para.add_run(sq.get('answer', ''))

    doc.add_paragraph()

def process_questions(doc, questions, q_type):
    """Process all questions of a given type."""
    for i, q in enumerate(questions, 1):
        if q_type == 'mcq':
            add_mcq_question(doc, i, q)
        elif q_type == 'true_false':
            add_true_false_question(doc, i, q)
        elif q_type == 'fill_in_blanks':
            add_fill_in_blanks_question(doc, i, q)
        elif q_type == 'match_the_following':
            add_match_question(doc, i, q)
        elif q_type == 'short_answer' or q_type == 'long_answer' or q_type == 'hots':
            add_short_answer_question(doc, i, q)
        elif q_type == 'textbook_qa':
            add_textbook_qa_question(doc, i, q)
        elif q_type == 'differentiate_between':
            add_differentiate_question(doc, i, q)
        elif q_type == 'give_reason':
            add_give_reason_question(doc, i, q)
        elif q_type == 'numericals':
            add_numerical_question(doc, i, q)
        elif q_type == 'assertion_reason':
            add_assertion_reason_question(doc, i, q)
        elif q_type == 'case_study':
            add_case_study_question(doc, i, q)
        elif q_type == 'name_the_following' or q_type == 'tricky_questions':
            add_short_answer_question(doc, i, q)
        else:
            # Generic handler
            add_textbook_qa_question(doc, i, q)

def create_chapter_document(subject, chapter, chapter_path):
    """Create a Word document for a chapter."""
    doc = create_document()
    add_chapter_title(doc, subject, chapter)

    # Find all question files
    question_files = [f for f in os.listdir(chapter_path) if f.endswith('.json') and f != 'sections.json']

    # Sort by preferred order
    preferred_order = ['textbook_qa', 'mcq', 'fill_in_blanks', 'true_false', 'match_the_following',
                       'short_answer', 'long_answer', 'assertion_reason', 'give_reason',
                       'differentiate_between', 'numericals', 'hots', 'case_study',
                       'name_the_following', 'tricky_questions']

    def sort_key(filename):
        name = filename.replace('.json', '')
        if name in preferred_order:
            return preferred_order.index(name)
        return len(preferred_order)

    question_files.sort(key=sort_key)

    for qfile in question_files:
        q_type = qfile.replace('.json', '')
        q_path = os.path.join(chapter_path, qfile)

        try:
            with open(q_path, 'r') as f:
                data = json.load(f)

            questions = data.get('questions', [])
            if not questions:
                continue

            # Add section title
            type_name = QUESTION_TYPE_NAMES.get(q_type, q_type.replace('_', ' ').title())
            add_question_type_title(doc, f"{type_name} ({len(questions)} questions)")

            # Add questions
            process_questions(doc, questions, q_type)

            # Add page break between sections
            doc.add_page_break()

        except Exception as e:
            print(f"Error processing {qfile}: {e}")

    return doc

def main():
    """Main function to create all Word documents."""
    # Create output directory
    os.makedirs(OUTPUT_PATH, exist_ok=True)

    # Get all subjects
    subjects_file = os.path.join(BASE_PATH, 'subjects.json')
    with open(subjects_file, 'r') as f:
        subjects_data = json.load(f)

    for subject in subjects_data['subjects']:
        subject_name = subject['name']
        subject_folder = subject_name.replace(' ', '_')
        subject_path = os.path.join(BASE_PATH, subject_folder)

        if not os.path.exists(subject_path):
            print(f"Subject folder not found: {subject_path}")
            continue

        # Get chapters
        chapters_file = os.path.join(subject_path, 'chapters.json')
        if not os.path.exists(chapters_file):
            print(f"Chapters file not found: {chapters_file}")
            continue

        with open(chapters_file, 'r') as f:
            chapters_data = json.load(f)

        for chapter in chapters_data['chapters']:
            chapter_name = chapter['name']
            # Try both with spaces and with underscores
            chapter_folder = chapter_name.replace(' ', '_')
            chapter_path = os.path.join(subject_path, chapter_folder)

            if not os.path.exists(chapter_path):
                # Try with spaces
                chapter_path = os.path.join(subject_path, chapter_name)

            if not os.path.exists(chapter_path):
                print(f"Chapter folder not found: {chapter_path}")
                continue

            print(f"Creating document for {subject_name} - {chapter_name}...")

            try:
                doc = create_chapter_document(subject_name, chapter_name, chapter_path)

                # Save document
                safe_filename = f"{subject_name}_{chapter_name}".replace(' ', '_').replace('/', '_')
                output_file = os.path.join(OUTPUT_PATH, f"{safe_filename}.docx")
                doc.save(output_file)
                print(f"  Saved: {output_file}")

            except Exception as e:
                print(f"  Error creating document: {e}")

    print(f"\nAll documents saved to: {OUTPUT_PATH}")

if __name__ == '__main__':
    main()
