#!/usr/bin/env python3
"""Generate a Word document from CS sample paper Python solutions JSON."""

import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Wine color
WINE = RGBColor(0x72, 0x2F, 0x37)
DARK_WINE = RGBColor(0x5B, 0x21, 0x2B)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Load exam questions JSON
JSON_PATH = 'questions_cache/Computer_Science/Python Programming/exam_questions.json'
with open(JSON_PATH, 'r') as f:
    data = json.load(f)

questions = data['questions']

# Group questions by source_section paper
papers = {}
for q in questions:
    source = q.get('source_section', 'Unknown')
    # Extract paper identifier (e.g., "Paper 1" from "Paper 1 — Predict the Output")
    paper_key = source.split(' — ')[0] if ' — ' in source else source
    section_name = source.split(' — ')[1] if ' — ' in source else 'Questions'

    if paper_key not in papers:
        papers[paper_key] = {}
    if section_name not in papers[paper_key]:
        papers[paper_key][section_name] = []
    papers[paper_key][section_name].append(q)

# Paper metadata
PAPER_INFO = {
    'Paper 1': 'Term I Examinations — July 2023 (80 marks)',
    'Paper 2': 'Term II Examinations — November 2023 (80 marks)',
    'Paper 4': 'Term I Examinations — July 2024 (80 marks)',
    'Paper 5': 'Term II Examinations — Nov/Dec 2024 (80 marks)',
    'Paper 6': 'Term I First Round Test — June 2025 (40 marks)',
    'Paper 7': 'Term I Second Round Test — July 2025 (40 marks)',
    'Paper 8': 'Term II Examinations — Nov/Dec 2025 (80 marks)',
}

# Sort papers in order
paper_order = ['Paper 1', 'Paper 2', 'Paper 4', 'Paper 5', 'Paper 6', 'Paper 7', 'Paper 8']


def add_wine_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = WINE
    return h


def add_section_heading(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = DARK_WINE
    return h


def add_question_answer(q, num):
    """Add a question-answer pair to the document."""
    # Question
    p = doc.add_paragraph()
    run = p.add_run(f'Q{num}. ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = WINE

    q_text = q['question']
    run = p.add_run(q_text)
    run.font.size = Pt(11)

    # Answer
    p = doc.add_paragraph()
    run = p.add_run('Answer: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_WINE

    a_text = q['answer']
    run = p.add_run(a_text)
    run.font.size = Pt(11)
    # Use monospace for code answers
    if '\n' in a_text or 'print(' in a_text or 'for ' in a_text or 'import ' in a_text or 'if ' in a_text:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    # Explanation (if present)
    if q.get('explanation'):
        p = doc.add_paragraph()
        run = p.add_run('Explanation: ')
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run = p.add_run(q['explanation'])
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Spacer
    doc.add_paragraph()


# ===== TITLE =====
title = doc.add_heading('Computer Science — Python Sample Paper Solutions', level=0)
for run in title.runs:
    run.font.color.rgb = WINE

subtitle = doc.add_paragraph()
run = subtitle.add_run('All 7 Papers: Python Questions Only (157 Questions)')
run.font.size = Pt(13)
run.font.color.rgb = DARK_WINE
run.italic = True

p = doc.add_paragraph()
run = p.add_run('Class 7ABC, Sishya School, Chennai')
run.font.size = Pt(11)
doc.add_paragraph()

# Generate each paper
for idx, paper_key in enumerate(paper_order):
    if paper_key not in papers:
        continue

    # Paper heading
    info = PAPER_INFO.get(paper_key, '')
    add_wine_heading(f'{paper_key}: {info}')

    sections = papers[paper_key]
    q_num = 1

    for section_name, section_questions in sections.items():
        add_section_heading(section_name)
        for q in section_questions:
            add_question_answer(q, q_num)
            q_num += 1

    # Page break between papers (except last)
    if idx < len(paper_order) - 1:
        doc.add_page_break()

# Save
output_path = '/Users/ramsabode/vibe-coding/edu_ani/subjects_class7c/Computer Science/sample papers/CSC_Sample_Papers_Python_Solutions.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
print(f'Total questions: {len(questions)}')
