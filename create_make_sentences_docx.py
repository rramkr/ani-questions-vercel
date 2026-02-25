import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Wine color
WINE = RGBColor(0x72, 0x2F, 0x37)
DARK_WINE = RGBColor(0x5B, 0x21, 0x2B)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Base path for question cache
BASE_PATH = '/Users/ramsabode/vibe-coding/ani-questions-vercel/questions_cache/English_E2'

# Chapter config: (display_name, folder_name)
CHAPTERS = [
    ('1. Daffodils — William Wordsworth', 'Daffodils'),
    ('2. The Picture Frame — Nicholas Horsburgh', 'The Picture Frame'),
    ('3. Shree Krishna Eating House — Anita Desai', 'Shree Krishna Eating House'),
    ('4. The Listeners — Walter de la Mare', 'The Listeners'),
]

# Helper functions
def add_wine_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = WINE
    return h

def add_wine_subheading(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = WINE
    return h

def add_section_label(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = DARK_WINE
    return h

def add_word_table(words_data):
    """words_data is list of (word, meaning, sentence)"""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(['Word', 'Meaning', 'Sentence']):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WINE
        run.font.size = Pt(11)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(10)

    # Data rows
    for word, meaning, sentence in words_data:
        row_cells = table.add_row().cells
        # Word in bold
        p = row_cells[0].paragraphs[0]
        run = p.add_run(word)
        run.bold = True
        run.font.size = Pt(10.5)
        # Meaning
        p = row_cells[1].paragraphs[0]
        run = p.add_run(meaning)
        run.font.size = Pt(10.5)
        # Sentence
        p = row_cells[2].paragraphs[0]
        run = p.add_run(sentence)
        run.font.size = Pt(10.5)

    doc.add_paragraph()  # spacing

# ===== TITLE =====
title = doc.add_heading('English E2 — Make Sentences', level=0)
for run in title.runs:
    run.font.color.rgb = WINE

subtitle = doc.add_paragraph()
run = subtitle.add_run('All Chapters: Words to Know & Additional Vocabulary')
run.font.size = Pt(13)
run.font.color.rgb = DARK_WINE
run.italic = True
doc.add_paragraph()

# Words to Know counts per chapter (textbook words come first in the JSON)
# These are the number of "Words to Know" entries at the start of each chapter's JSON
WORDS_TO_KNOW_COUNT = {
    'Daffodils': 10,
    'The Picture Frame': 7,
    'Shree Krishna Eating House': 6,
    'The Listeners': 14,
}

# Generate each chapter
for idx, (display_name, folder_name) in enumerate(CHAPTERS):
    add_wine_heading(display_name)

    # Load the make_sentences.json
    json_path = f'{BASE_PATH}/{folder_name}/make_sentences.json'
    with open(json_path, 'r') as f:
        data = json.load(f)

    questions = data['questions']
    wtk_count = WORDS_TO_KNOW_COUNT[folder_name]

    # Split into Words to Know and Additional
    wtk_words = questions[:wtk_count]
    additional_words = questions[wtk_count:]

    # Words to Know section
    add_section_label('Words to Know (Textbook)')
    add_word_table([
        (q['word'], q['meaning'], q['sample_sentence'])
        for q in wtk_words
    ])

    # Additional Vocabulary section (if any)
    if additional_words:
        add_section_label('Additional Vocabulary')
        add_word_table([
            (q['word'], q['meaning'], q['sample_sentence'])
            for q in additional_words
        ])

    # Page break between chapters (except last)
    if idx < len(CHAPTERS) - 1:
        doc.add_page_break()

# Save
output_path = '/Users/ramsabode/vibe-coding/edu_ani/subjects_class7c/E2/sample paper/E2_Make_Sentences_All_Chapters.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
